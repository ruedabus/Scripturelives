#!/usr/bin/env python3
"""
Scripture Lives — DOCX Study Export
Reads JSON from stdin, writes a formatted Word document to stdout (binary).
"""

import sys
import json
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

def set_heading_color(paragraph, rgb: tuple):
    """Set colour on all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)


def add_horizontal_rule(doc):
    """Add a thin horizontal line between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D4A853')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_label_value(doc, label: str, value: str):
    """Add a paragraph like  Label: value  with bold label."""
    if not value or not value.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)  # stone-500
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)


# ──────────────────────────────────────────────────────────────
# Document builder
# ──────────────────────────────────────────────────────────────

def build_document(data: dict) -> bytes:
    doc = Document()

    # ── Page margins ───────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ── Title ──────────────────────────────────────────────────
    title = doc.add_heading("Scripture Lives", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_color(title, (0xD4, 0xA8, 0x53))  # amber

    subtitle = doc.add_paragraph("Study Summary")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.runs[0] if subtitle.runs else subtitle.add_run("Study Summary")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

    date_str = datetime.now().strftime("%B %d, %Y  %-I:%M %p")
    date_p = doc.add_paragraph(f"Exported: {date_str}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_p.runs:
        date_p.runs[0].font.size = Pt(9)
        date_p.runs[0].font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)

    doc.add_paragraph()  # spacer

    # ──────────────────────────────────────────────────────────
    # Section 1 — Bookmarked Places
    # ──────────────────────────────────────────────────────────
    bookmarks = data.get("bookmarks", [])   # list of place objects
    notes     = data.get("notes", {})       # { placeName: noteText }

    add_horizontal_rule(doc)
    h1 = doc.add_heading("Bookmarked Places", level=1)
    set_heading_color(h1, (0xD4, 0xA8, 0x53))

    if not bookmarks:
        p = doc.add_paragraph("No bookmarked places saved.")
        p.runs[0].font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
    else:
        for i, item in enumerate(bookmarks, 1):
            place  = item.get("place", {})
            name   = item.get("name", "Unknown Place")
            ref    = item.get("sourceReference", "")
            note   = notes.get(name, item.get("note", ""))

            # Place name as H2
            h2 = doc.add_heading(f"{i}. {name}", level=2)
            set_heading_color(h2, (0x1C, 0x17, 0x10))  # dark

            add_label_value(doc, "Era",           place.get("era", ""))
            add_label_value(doc, "Source Verse",  ref)
            add_label_value(doc, "Summary",       place.get("description", ""))
            add_label_value(doc, "Biblical Significance", place.get("biblicalSignificance", ""))

            # Study Note
            note_p = doc.add_paragraph()
            note_p.paragraph_format.space_before = Pt(4)
            note_label = note_p.add_run("My Note: ")
            note_label.bold = True
            note_label.font.size = Pt(10)
            note_label.font.color.rgb = RGBColor(0xD4, 0xA8, 0x53)
            note_text = note_p.add_run(note if note and note.strip() else "No note saved.")
            note_text.font.size = Pt(10)
            note_text.italic = not (note and note.strip())
            note_text.font.color.rgb = (
                RGBColor(0x1C, 0x17, 0x10) if (note and note.strip())
                else RGBColor(0xA8, 0xA2, 0x9E)
            )

            if i < len(bookmarks):
                doc.add_paragraph()

    # ──────────────────────────────────────────────────────────
    # Section 2 — Study Sessions
    # ──────────────────────────────────────────────────────────
    sessions = data.get("sessions", [])

    if sessions:
        doc.add_paragraph()
        add_horizontal_rule(doc)
        h1s = doc.add_heading("Study Sessions", level=1)
        set_heading_color(h1s, (0xD4, 0xA8, 0x53))

        for j, session in enumerate(sessions, 1):
            created = session.get("createdAt", "")
            try:
                dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
                created_fmt = dt.strftime("%B %d, %Y  %-I:%M %p")
            except Exception:
                created_fmt = created

            h2s = doc.add_heading(
                f"{j}. {session.get('name', 'Unnamed Session')}",
                level=2
            )
            set_heading_color(h2s, (0x1C, 0x17, 0x10))

            add_label_value(doc, "Saved on", created_fmt)

            sess_bookmarks = session.get("bookmarks", [])
            sess_notes     = session.get("notes", {})

            if sess_bookmarks:
                bm_p = doc.add_paragraph()
                bm_p.paragraph_format.space_before = Pt(4)
                bm_label = bm_p.add_run("Bookmarks: ")
                bm_label.bold = True
                bm_label.font.size = Pt(10)
                bm_text = bm_p.add_run(", ".join(sess_bookmarks))
                bm_text.font.size = Pt(10)
            else:
                p = doc.add_paragraph("No places bookmarked in this session.")
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
                    p.runs[0].font.size = Pt(10)

            # Session notes
            if sess_notes:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                r = p.add_run("Session Notes:")
                r.bold = True
                r.font.size = Pt(10)
                for place_name, note_text in sess_notes.items():
                    if note_text and note_text.strip():
                        np = doc.add_paragraph(style="List Bullet")
                        np.paragraph_format.left_indent = Cm(0.5)
                        run_n = np.add_run(f"{place_name}: ")
                        run_n.bold = True
                        run_n.font.size = Pt(10)
                        run_t = np.add_run(note_text)
                        run_t.font.size = Pt(10)

            if j < len(sessions):
                doc.add_paragraph()

    # ──────────────────────────────────────────────────────────
    # Footer note
    # ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_horizontal_rule(doc)
    footer_p = doc.add_paragraph("Generated by Scripture Lives  ·  Study deeper. Explore further.")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_p.runs:
        footer_p.runs[0].font.size = Pt(8)
        footer_p.runs[0].font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
        footer_p.runs[0].italic = True

    # ── Serialize to bytes ─────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────
# Entry point — read JSON from stdin, write DOCX to stdout
# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    raw = sys.stdin.buffer.read()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"Invalid JSON: {e}\n")
        sys.exit(1)

    doc_bytes = build_document(data)
    sys.stdout.buffer.write(doc_bytes)
